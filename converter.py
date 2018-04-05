import os
from docx import Document

doc = Document()

def convert(dir, mode='flat', title=None, include=None, exclude=None, encoding='utf-8'):
  print('copy from diretory: ' + dir)

  if title is not None:
    doc.add_heading(title, 1)
  
  if include is not None:
    inc=include.split('|')
  else:
    inc=None
  
  if exclude is not None:
    exc=exclude.split('|')
  else:
    exc=None

  if mode == 'flat':
    walkflat(dir, inc, exc, encoding)
  elif mode == 'deep':
    walkdeep(dir, 2, inc, exc, encoding)
  else:
    print('mode is invaild')


def walkflat(dir, inc, exc, encoding):
  currentdir = ''
  for root, dirs, files in os.walk(dir,False):
    for file in files:
      if (inc is None or os.path.splitext(file)[1][1:] in inc) and (exc is None or os.path.splitext(file)[1][1:] not in exc):
        filepath = os.path.join(root,file).replace('\\','/')
        with open(filepath,encoding=encoding) as f:
          content = f.read()
          thisdir = filepath[len(dir)+1:filepath.rfind('/')]
          if currentdir != thisdir:
            currentdir = thisdir
            doc.add_heading(thisdir, 2)
            print('into directory '+thisdir)
          doc.add_heading(filepath[filepath.rfind('/')+1:], 3)
          doc.add_paragraph(content)
          doc.add_page_break()
          print('copying '+filepath[filepath.rfind('/')+1:])

def walkdeep(root, level, inc, exc, encoding):
  for file in os.listdir(root):
    filepath = os.path.join(root,file).replace('\\','/')
    if os.path.isfile(filepath):
      if (inc is None or os.path.splitext(file)[1][1:] in inc) and (exc is None or os.path.splitext(file)[1][1:] not in exc):
        with open(filepath,encoding=encoding) as f:
          content = f.read()
          doc.add_heading(filepath[filepath.rfind('/')+1:], level)
          doc.add_paragraph(content)
          doc.add_page_break()
          print('copying '+filepath[filepath.rfind('/')+1:])
    else:
      doc.add_heading(file, level)
      print('into directory '+file)
      walkdeep(filepath, level+1, inc, exc, encoding)